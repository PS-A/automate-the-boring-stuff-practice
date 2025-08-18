# Custom Invitations as Word Documents

import docx


def create_invitations(guestlist):

    doc = docx.Document()
    with open(guestlist, "r") as file:
        for line in file:
            doc.add_paragraph(f"Hello {line}!", "Title")
            doc.add_heading("You are invited to a party!", 3)
            doc.add_heading("Welcome!", 4)
            doc.add_page_break()
    doc.save("guests.docx")


create_invitations("guests.txt")
