# Project: Combining Select Pages from Many PDFS
"""
#! python3
import PyPDF2, os
   # Get all the PDF filenames.
pdfFiles = []
for filename in os.listdir('.'):
    if filename.endswith(".pdf"):
        pdfFiles.append(filename)
pdfFiles.sort(key = str.lower)

pdfWriter = PyPDF2.PdfFileWriter()

# Loop through all the PDF files.
for filename in pdfFiles:
    pdfFileObj = open(filename, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    # Loop through all the pages (except the first) and add them.
    for pageNum in range(1, pdfReader.numPages):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)
# Save the resulting PDF to a file.
pdfOutput = open('allminutes.pdf', 'wb')
pdfWriter.write(pdfOutput)
pdfOutput.close()
"""

# Read Docx
"""
#! python3

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
"""

# Convert Word To PDF
"""
# This script runs on Windows only, and you must have Word installed.
import win32com.client # install with "pip install pywin32"
import docx
wordFilename = 'test.docx'
pdfFilename = 'test.pdf'

doc = docx.Document()
# Code to create Word document goes here.
doc.save(wordFilename)

wdFormatPDF = 17 # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()
"""